import docx
from docx import Document
from docx import oxml
from docx.opc import constants
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches, Pt, RGBColor as Color


"""Returns a document object"""


def new_document():
    doc = Document()
    return doc


"""Takes doc object and title as parameters and adds Heading to doc"""


def add_title_to_doc(doc, title):
    doc.add_heading(title, 0)


"""Takes doc object, title and heading size as parameters and adds subheadings to doc based on heading size"""


def add_sub_heading_to_doc(doc, heading, heading_size=1):
    doc.add_heading(heading, heading_size)


"""Helps to add pictures to doc. Takes doc object and path of image along with size to add image."""


def add_picture_to_doc(doc, path, width=None, height=None):
    doc.add_picture(path, width=Inches(width), height=Inches(height))


"""Helps to add paragraph to doc and returns para object.Takes doc object and font format options as input parameters"""


def add_paragraph_to_doc(doc, text=None, style=None, font_name='Calibri', font_size=10, font_color=(0, 0, 0),
                         bold=False, italic=False, underline=False, space_before=0, left_indent=1):
    if style:
        style = style.lower()
        para_dict = {'bullet': "List Bullet", 'number': "List Number", 'quote': "Intense Quote"}
        para = doc.add_paragraph(text, style=para_dict[style])
    else:
        para = doc.add_paragraph(text)
    para.style.font.size = Pt(font_size)
    para.style.font.name = font_name
    para.style.font.color.rgb = Color(font_color[0], font_color[1], font_color[2])
    para.style.font.bold = bold
    para.style.font.italic = italic
    para.style.font.underline = underline
    para_format = para.paragraph_format
    para_format.space_before = Inches(space_before)
    para_format.left_indent = Inches(left_indent)
    return para


"""Helps to continue existing para. Takes para object and text formatting options."""


def continue_para_in_doc(para, text, bold=False, italic=False, underline=False):
    run = para.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline


"""Helps to add hyperlink to doc. Takes para object, text and url as input parameters."""


def add_hyperlink_to_doc(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


"""Helps to add table to doc. 
Takes doc object, number of rows and columns as input parameters and returns table object"""


def add_table_to_doc(doc, rows, columns):
    table = doc.add_table(rows=rows, cols=columns)
    return table


"""Helps to edit created table.Takes table object and text to add in it along with style"""


def edit_table_in_doc(table, text, row, col, style="Colorful List"):
    cell = table.rows[row].cells
    cell[col].text = text
    table.style = style
